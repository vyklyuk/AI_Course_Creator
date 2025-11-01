"""
Case Study Lecture Creation Agent — FINAL (EN, bold headings, auto-retry, Word tables)
- External OpenAI client is injected.
- Generates narrative lecture text in a Case Study learning format.

Parameters per run:
  1) course_title        (str)
  2) module_title        (str)
  3) lecture_title       (str)
  4) case_study_desc     (str)   # concise scenario description (facts + context)
  5) learning_outcomes   (List[str] | str)
  6) minutes             (int)   # 8000 words/hour budget for the full lecture text
  7) save_path           (Optional[str]; .docx enforced)

Requires:
  pip install openai python-docx
"""

from __future__ import annotations
from typing import Optional, Tuple, List, Union
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.shared import Pt
import re

# ------------------------- Style helpers -------------------------

APPROVED_TRANSITIONS: List[str] = [
    "Another angle to examine is...",
    "Building on the prior step, let's analyze...",
    "With these facts in mind, consider...",
    "From here, we can compare...",
    "To pressure-test this, let's explore...",
    "Shifting focus, look at...",
    "This brings us to a decision point...",
    "Keeping that in view, evaluate...",
    "Now, let's break down the evidence for...",
    "Moving forward, quantify...",
]

DIFFICULTY_PROFILES = {
    "Beginner": (
        "- Audience: newcomers; avoid jargon or define on first use.\n"
        "- Scaffolding: heavy. Provide definitions, mini-examples, check-for-understanding.\n"
        "- Math: minimal. One very simple calculation.\n"
        "- Data Exhibits: 1–2 small tables (≤5 rows). Keep numbers round.\n"
        "- Guiding Questions: 0–3 very simple [Apply] items, or omit entirely.\n"  
        "- Sentences: keep very short; explain acronyms.\n"
    ),
    "Intermediate": (
        "- Audience: practitioners; moderate jargon with quick clarifications.\n"
        "- Scaffolding: moderate. Move from facts to patterns to principles.\n"
        "- Math: 1–2 quick calculations; simple ratios/percents.\n"
        "- Data Exhibits: 2–3 tables (≤8 rows) with plausible figures.\n"
        "- Guiding Questions: 2–5 with mostly [Apply] and a few [Analyze].\n"  # 
    ),
    "Advanced": (
        "- Audience: experts; domain terminology without basic definitions.\n"
        "- Scaffolding: light. Emphasize trade-offs and assumptions.\n"
        "- Math: multiple calculations or a quick sensitivity check.\n"
        "- Data Exhibits: 2–3 denser tables (≤10 rows) enabling comparisons.\n"
        "- Guiding Questions: 4–6 with [Analyze]/[Evaluate] emphasis; include counterfactuals.\n"  
        "- Include edge cases and limitations explicitly.\n"
    ),
    "Executive": (
        "- Audience: non-technical decision-makers.\n"
        "- Focus: implications, risk, ROI, and decisions; minimal math.\n"
        "- Data Exhibits: 1–2 summary tables; highlight deltas and thresholds.\n"
        "- Guiding Questions: 2–4 oriented to trade-offs and action.\n"  
        "- Keep explanations strategic and concise; avoid low-level details.\n"
    ),
}

def _normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())

# ------------------------- Prompt templates (EN only) -------------------------

def build_difficulty_block(level: str) -> str:
    lvl = (level or "Intermediate").lower().strip()
    if lvl not in DIFFICULTY_PROFILES:
        lvl = "Intermediate"
    return (
        f"\nDIFFICULTY PROFILE (active: {lvl}):\n"
        f"{DIFFICULTY_PROFILES[lvl]}"
        "Apply this profile consistently across tone, examples, math depth, data exhibits, and question difficulty.\n"
    )


SYSTEM_TEMPLATE_EN = (
    "You are a senior instructional designer and subject-matter writing coach.\n"
    "Your job is to produce a clear, engaging lecture written as a Case Study for higher education.\n\n"
    "STRICT STYLE RULES:\n"
    "1) Pedagogy: Case Method. Students encounter a realistic scenario, analyze evidence, debate options, and justify decisions.\n"
    "2) Voice: instructive but conversational; short sentences; average <= 18 words.\n"
    "3) Scaffolding: move from concrete facts to abstract principles; include prompts that target Bloom's levels (apply, analyze, evaluate, create).\n"
    "4) Early-stage cognitive load: avoid overwhelming beginners—keep pre-analysis questions minimal or omit them.\n"  # NEW
    "5) Rhetorical questions: use sparingly to provoke thinking. No fluff.\n"
    "6) Data realism: include small, plausible data exhibits. Format each exhibit as a pipe-separated table with a header row, e.g. Metric|Q1|Q2. Do not use Markdown borders.\n"
    "7) Academic integrity: avoid copyrighted text and personally identifiable data; create synthetic but realistic details.\n"
    "8) ASCII only quotes and ellipses: \" \" and ...\n\n"
    "OUTPUT POLICY:\n"
    "- Plain text only (no JSON, no markdown tables, no code fences).\n"
    "- Follow the output template exactly. Keep sections in this order and include all separators.\n"
    "- Target total ~{target_words} words across the whole lecture (±10%).\n"
    "- Write in English.\n"
)

OUTPUT_TEMPLATE_EN = (
    "Title: {lecture_title}\n"
    "Course: {course_title}\n"
    "Module: {module_title}\n"
    "Duration: {minutes} minutes\n"
    "---\n"
    "Learning Outcomes:\n"
    "- {los}\n"
    "---\n"
    "Case Overview:\n"
    "Write a tight narrative (120–200 words) that sets the scene using the scenario below. Name key stakeholders and the central tension.\n"
    "Scenario input: {case_study_desc}\n"
    "---\n"
    "Context & Timeline:\n"
    "Provide 5–8 bullet sentences that anchor the case in time, place, constraints, and assumptions.\n"
    "---\n"
    "Data Exhibits:\n"
    "Create 2–3 small exhibits as pipe-separated tables (header in first row). Keep each exhibit under 10 rows. Example line: Metric|Q1|Q2|Note\n"
    "---\n"
    "{guiding_questions_block}"
    "Analysis Walkthrough:\n"
    "Write a stepwise reasoning path using short paragraphs. Each step should start with one of these transitions (vary; do not repeat consecutively): {approved_transitions}\n"
    "Include at least one quick calculation and one assumption check.\n"
    "---\n"
    "Decision Points:\n"
    "List 2–4 options. For each, provide: expected impact, key risk, and what evidence would confirm/kill it.\n"
    "---\n"
    "Instructor Notes (Model Answers):\n"
    "{instructor_notes_line}\n"
    "---\n"
    "Wrap-Up & Transfer:\n"
    "Summarize 4–6 key takeaways and show how to transfer the insights to new contexts.\n"
    "---\n"
    "Reflection Prompt:\n"
    "One short prompt that nudges students to connect the case to their own context.\n"
)

SECTION_HEADERS = {
    "Learning Outcomes:",
    "Case Overview:",
    "Context & Timeline:",
    "Data Exhibits:",
    "Guiding Questions (Before Analysis):",
    "Analysis Walkthrough:",
    "Decision Points:",
    "Instructor Notes (Model Answers):",
    "Wrap-Up & Transfer:",
    "Reflection Prompt:",
}



def _extract_headered_json(text: str) -> dict:
    """
    Build a JSON dict whose keys mirror HEADER_PREFIXES and SECTION_HEADERS.
    - HEADER_PREFIXES entries take the single-line value after "Label: ..."
    - SECTION_HEADERS entries accumulate lines after the header until the next header.
    Missing keys are included with "".
    """
    prefixes = list(HEADER_PREFIXES)
    sections = list(SECTION_HEADERS)
    result = {k: "" for k in prefixes}
    result.update({k: "" for k in sections})
    lines = [ln.rstrip("\n") for ln in text.splitlines()]

    # Prefix lines (single-line values)
    for ln in lines:
        s = ln.strip()
        for pref in prefixes:
            if s.startswith(pref):
                val = s[len(pref):].strip()
                if val.startswith(":"):
                    val = val[1:].strip()
                result[pref] = val
                break

    # Section blocks
    indices = []
    set_sections = set(sections)
    for idx, ln in enumerate(lines):
        s = ln.strip()
        if s in set_sections:
            indices.append((idx, s))

    for i, (start_idx, header) in enumerate(indices):
        end_idx = indices[i+1][0] if i+1 < len(indices) else len(lines)
        block = lines[start_idx+1:end_idx]
        while block and not block[0].strip():
            block = block[1:]
        while block and not block[-1].strip():
            block = block[:-1]
        result[header] = "\n".join(block).strip()

    return result
HEADER_PREFIXES = ("Title:", "Course:", "Module:", "Duration:")

# ------------------------- Core helpers -------------------------

def calc_target_words(minutes: int) -> int:
    return round(minutes * (8000 / 60.0))

def build_system(minutes: int, difficulty: str = "Intermediate") -> str:
    base = SYSTEM_TEMPLATE_EN.format(target_words=calc_target_words(minutes))
    return base + build_difficulty_block(difficulty)


def build_user_prompt(
    course_title: str,
    module_title: str,
    lecture_title: str,
    case_study_desc: str,
    learning_outcomes: Union[List[str], str],
    minutes: int,
    difficulty: str = "Intermediate",
    guiding_questions_mode: str = "standard",   
) -> str:
    los = learning_outcomes
    if isinstance(los, list):
        los = "\n- ".join(los)
    else:
        los = "\n- " + "\n- ".join([p.strip() for p in re.split(r"[;\n]", los) if p.strip()])

    gq_block, notes_line = _build_guiding_questions_block(difficulty, guiding_questions_mode)

    core = OUTPUT_TEMPLATE_EN.format(
        lecture_title=lecture_title,
        course_title=course_title,
        module_title=module_title,
        minutes=minutes,
        los=los.strip(),
        case_study_desc=case_study_desc.strip(),
        approved_transitions=", ".join(APPROVED_TRANSITIONS),
        guiding_questions_block=gq_block,              
        instructor_notes_line=notes_line,              
    )
    return core


def _build_guiding_questions_block(difficulty: str, mode: str) -> Tuple[str, str]:
    """
    Returns (guiding_questions_block, instructor_notes_line).
    mode: "none" | "light" | "standard"
    """
    mode = (mode or "standard").strip().lower()
    # Підбираємо кількість і мітки залежно від режиму й складності
    if mode == "none":
        return ("", "Provide a recommended option with justification. Note typical student pitfalls and misconceptions.")  # no Qs

    if mode == "light":
        # Стислі інструкції для простих питань
        q_text = (
            "Guiding Questions (Before Analysis):\n"
            "Write 2–3 very simple questions focused on checking comprehension and immediate application. "
            "Prefer [Apply]; avoid multi-step, abstract, or compound questions.\n"
            "---\n"
        )
        notes = (
            "Provide concise model responses for the brief Guiding Questions and a recommended option with justification. "
            "Highlight common beginner misconceptions."
        )
        return (q_text, notes)

    # "standard"
    q_text = (
        "Guiding Questions (Before Analysis):\n"
        "Write 4–6 questions that move from facts to hypotheses. Tag each with [Apply], [Analyze], or [Evaluate]. "
        "Avoid compound questions; use one skill per question.\n"
        "---\n"
    )
    notes = "Provide concise model responses for the Guiding Questions and a recommended option with justification. Note typical student pitfalls and misconceptions."
    return (q_text, notes)


# ------------------------- Post processors -------------------------

_transitions_pattern = re.compile(
    r"^(Another angle to examine is\.|Building on the prior step, let's analyze\.|With these facts in mind, consider\.|From here, we can compare\.|To pressure-test this, let's explore\.|Shifting focus, look at\.|This brings us to a decision point\.|Keeping that in view, evaluate\.|Now, let's break down the evidence for\.|Moving forward, quantify\.)",
    re.I,
)

def fix_consecutive_transitions(text: str) -> str:
    lines = text.splitlines()
    prev = ""
    for i, ln in enumerate(lines):
        m = _transitions_pattern.match(ln.strip())
        if m:
            cur = m.group(0)
            if _normalize(cur) == _normalize(prev):
                for cand in APPROVED_TRANSITIONS:
                    if _normalize(cand) != _normalize(prev):
                        lines[i] = ln.replace(cur, cand, 1)
                        prev = cand
                        break
            else:
                prev = cur
    return "\n".join(lines)

# ------------------------- DOCX writer (tables + bold headers) -------------------------

def write_docx_with_formatting(text: str, path: Path) -> None:
    doc = Document()
    table_buffer: List[List[str]] = []

    def looks_like_table_line(line: str) -> bool:
        if "|" not in line:
            return False
        if set(line.strip()) <= set("-+| "):
            return False  # skip markdown borders
        parts = [c.strip() for c in line.strip().strip("|").split("|")]
        return sum(1 for c in parts if c) >= 2

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        max_cols = max(len(r) for r in table_buffer)
        table: Table = doc.add_table(rows=len(table_buffer), cols=max_cols)
        table.style = "Table Grid"
        for r_i, row in enumerate(table_buffer):
            for c_i in range(max_cols):
                txt = row[c_i] if c_i < len(row) else ""
                table.cell(r_i, c_i).text = txt
        table_buffer = []

    for raw in text.splitlines():
        line = raw.rstrip("\n")

        # Collect table rows
        if looks_like_table_line(line):
            parts = [c.strip() for c in line.strip().strip("|").split("|")]
            table_buffer.append(parts)
            continue

        # Non-table line -> flush any pending table first
        flush_table()

        p = doc.add_paragraph()
        s = line.strip()

        # Bold for section headers (exact match line)
        if s in SECTION_HEADERS:
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Bold for label prefixes Title/Course/Module/Duration (label bold, value normal)
        prefix = next((pref for pref in HEADER_PREFIXES if s.startswith(pref)), None)
        if prefix:
            # Keep original spacing and value
            if ":" in line:
                label, rest = line.split(":", 1)
                run_label = p.add_run(label + ":")
                run_label.bold = True
                run_label.font.size = Pt(12)
                if rest:
                    run_val = p.add_run(rest)
                    run_val.font.size = Pt(11)
            else:
                run_label = p.add_run(line)
                run_label.bold = True
                run_label.font.size = Pt(12)
            continue

        # Default paragraph
        run = p.add_run("" if not s else line)
        run.font.size = Pt(11)

    # Flush trailing table if any
    flush_table()

    doc.save(path)

# ------------------------- Agent -------------------------

class CaseStudyLectureAgent:
    """Agent that generates a case-study style lecture text. Uses an injected OpenAI client.

    Usage:
        from openai import OpenAI
        client = OpenAI(api_key="...")
        agent = CaseStudyLectureAgent(client, model="gpt-5-nano")
        text, path = agent.generate(...)
    """

    def __init__(self, client, model: str = "gpt-5-nano"):
        self.client = client
        self.model = model
        # store last generation payload for optional JSON export
        self._last_payload = None  # type: ignore


    def _call_model(self, system: str, user: str) -> str:
        resp = self.client.responses.create(
            model=self.model,
            input=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        )
        text = getattr(resp, "output_text", None)
        if not text:
            try:
                text = resp.output[0].content[0].text
            except Exception:
                text = str(resp)
        return text

    def generate(
        self,
        course_title: str,
        module_title: str,
        lecture_title: str,
        case_study_desc: str,
        learning_outcomes: Union[List[str], str],
        difficulty: str = "Intermediate",
        minutes: int = 30,
        save_path: Optional[str] = None,
        guiding_questions_mode: Optional[str] = None,
        verbose: Optional[bool] = True,
        exam: bool = False,
        ) -> Tuple[str, Optional[str]]:
        system = build_system(minutes, difficulty=difficulty)
        user_prompt = build_user_prompt(
            course_title=course_title,
            module_title=module_title,
            lecture_title=lecture_title,
            case_study_desc=case_study_desc,
            learning_outcomes=learning_outcomes,
            minutes=minutes,
            difficulty=difficulty,
            guiding_questions_mode=guiding_questions_mode,
        )

        # First call
        if verbose:
            print("System:\n", system, "\nUser:\n", user_prompt)
            
        text = self._call_model(system, user_prompt)

        # Auto-retry if below target words (with safety cap)
        target = calc_target_words(minutes)
        words = len(re.findall(r"\b\w+\b", text))
        retries = 0
        while words < target * 0.9 and retries < 2:
            addendum = (
                f"The previous draft was {words} words. Expand the lecture to approximately {target} words (±10%). "
                "Keep the same sections and order; do not add new sections. "
                "Elaborate with deeper explanations, examples, short calculations, and clear justifications. "
                "Maintain the Case Study pedagogy and transitions."
            )
            text = self._call_model(system, user_prompt + "\n" + addendum)
            words = len(re.findall(r"\b\w+\b", text))  # update word count after retry
            retries += 1
        text = fix_consecutive_transitions(text)

        out_path: Optional[str] = None
        if save_path:
            p = Path(save_path)
            if p.suffix.lower() != ".docx":
                p = p.with_suffix(".docx")
            write_docx_with_formatting(text, p)
            out_path = str(p)

            if exam:
                exam_path = self._write_exam_docx(text, p)
                if verbose:
                    print(f"Exam version saved: {exam_path}")

        # cache last payload for JSON export
        self._last_payload = {
            "text": text,
            "course_title": course_title,
            "module_title": module_title,
            "lecture_title": lecture_title,
            "case_study_desc": case_study_desc,
            "learning_outcomes": learning_outcomes,
            "minutes": minutes,
            "difficulty": difficulty,
            "guiding_questions_mode": guiding_questions_mode,
            "docx_path": out_path,
        }

        return text, out_path


    
    def save_to_json(self, json_path: str | None = None, payload: dict | None = None) -> str:
            """
            Save the last generated lecture to JSON with keys matching HEADER_PREFIXES and SECTION_HEADERS.
            If json_path is omitted, it mirrors the last .docx filename with .json extension.
            """
            import json
            from pathlib import Path as _Path
    
            data = payload if payload is not None else getattr(self, "_last_payload", None)
            if data is None:
                raise ValueError("No data to save. Call `generate(...)` first or pass `payload`.")
    
            # Determine target path
            if not json_path or not str(json_path).strip():
                docx_path = data.get("docx_path") if isinstance(data, dict) else None
                if not docx_path:
                    raise ValueError("json_path not provided and no docx_path cached from generate().")
                p = _Path(docx_path).with_suffix(".json")
            else:
                p = _Path(json_path)
                if p.suffix.lower() != ".json":
                    p = p.with_suffix(".json")
    
            # Build structured JSON from text
            text = data.get("text") if isinstance(data, dict) else None
            if not isinstance(text, str) or not text.strip():
                raise ValueError("Cached payload has no text to serialize.")
            structured = _extract_headered_json(text)
    
            p.write_text(json.dumps(structured, ensure_ascii=False, indent=2), encoding="utf-8")
            return str(p)

    def _write_exam_docx(self, text: str, original_path: Path) -> str:
        """
        Creates content_exam.docx with only 4 sections:
        - Learning Outcomes
        - Case Overview
        - Context & Timeline
        - Data Exhibits
        """
        exam_path = original_path.with_name("content_exam.docx")
        doc = Document()
        table_buffer: List[List[str]] = []
    
        # Ті ж самі 4 заголовки, що потрібні
        allowed_sections = {
            "Learning Outcomes:",
            "Case Overview:",
            "Context & Timeline:",
            "Data Exhibits:",
        }
    
        def looks_like_table_line(line: str) -> bool:
            if "|" not in line:
                return False
            if set(line.strip()) <= set("-+| "):
                return False
            parts = [c.strip() for c in line.strip().strip("|").split("|")]
            return sum(1 for c in parts if c) >= 2
    
        def flush_table():
            nonlocal table_buffer
            if not table_buffer:
                return
            max_cols = max(len(r) for r in table_buffer)
            table: Table = doc.add_table(rows=len(table_buffer), cols=max_cols)
            table.style = "Table Grid"
            for r_i, row in enumerate(table_buffer):
                for c_i in range(max_cols):
                    txt = row[c_i] if c_i < len(row) else ""
                    table.cell(r_i, c_i).text = txt
            table_buffer = []
    
        current_section = None
        in_allowed_section = False
    
        for raw in text.splitlines():
            line = raw.rstrip("\n")
            s = line.strip()
    
            # Визначаємо, чи це заголовок розділу
            if s in SECTION_HEADERS:
                flush_table()  # скидаємо таблицю перед новим розділом
                current_section = s
                in_allowed_section = s in allowed_sections
    
                if in_allowed_section:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(12)
                continue
    
            # Якщо ми не в дозволеному розділі — пропускаємо
            if not in_allowed_section:
                continue
    
            # Обробка таблиці
            if looks_like_table_line(line):
                parts = [c.strip() for c in line.strip().strip("|").split("|")]
                table_buffer.append(parts)
                continue
    
            # Скидаємо таблицю перед новим абзацом
            flush_table()
    
            # Додаємо звичайний текст
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
    
        flush_table()  # остання таблиця
        doc.save(exam_path)
        return str(exam_path)
# ------------------------- Minimal example -------------------------
if __name__ == "__main__":
    try:
        from openai import OpenAI
        client = OpenAI()  # expects env var OPENAI_API_KEY
        agent = CaseStudyLectureAgent(client, model="gpt-5-nano")
        text, path = agent.generate(
            course_title="Data Ethics and Governance",
            module_title="Responsible AI in Practice",
            lecture_title="Bias Escalation in Credit Scoring",
            case_study_desc=(
                "A mid-size bank pilots a new ML credit model. Approval rates rise, "
                "but complaints spike among self-employed applicants. Regulators request a review."
            ),
            learning_outcomes=[
                "Diagnose dataset and process sources of bias",
                "Quantify impacts using simple fairness metrics",
                "Formulate mitigations and trade-offs",
                "Argue a decision with evidence",
            ],
            minutes=15,
            save_path="case_lecture.docx",
            difficulty="Beginner",
            guiding_questions_mode="none",
        )
        print(path or "<no file saved>")
    except Exception as e:
        print("Example run skipped:", e)
