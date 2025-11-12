import os
import json
import re
import random
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

try:
    import docx as docx_reader
except Exception:
    docx_reader = None


class CaseExerciseAgent:
    def __init__(self, client, model: str = "gpt-5-nano"):
        self.client = client
        self.model = model

    def generate(
        self,
        case_file: str,
        n_questions: int = 3,
        n_answers: int = 10,
        n_correct: int = 3,
        shuffle_positions: bool = True,
        seed: int | None = None,
        questions: Optional[List[str]] = None,
        verbose: Optional[bool] = True,
    ) -> Tuple[str, str, str]:
        self._validate_counts(n_answers, n_correct)
        case_text = self._load_case_studies(case_file)

        use_provided = questions is not None and len(questions) > 0
        effective_n_questions = len(questions) if use_provided else n_questions
        if effective_n_questions <= 0:
            raise ValueError("Кількість запитань має бути > 0.")

        # --- Генерація даних ---
        data = self._call_model_structured(
            case_text=case_text,
            n_questions=effective_n_questions,
            n_answers=n_answers,
            n_correct=n_correct,
            provided_questions=questions if use_provided else None,
        )

        raw_exercises = data.get("exercises", [])[:effective_n_questions]

        # --- Нормалізація + перемішування ---
        rng = random.Random(seed) if seed is not None else random.Random()
        exercises = []  # ← Оголошуємо ТУТ
        for ex in raw_exercises:
            ex_norm = self._normalize_exercise(ex, n_answers)
            self._enforce_counts(ex_norm, n_answers, n_correct)
            if shuffle_positions:
                ex_norm = self._shuffle_and_relabel_options(ex_norm, rng)
            exercises.append(ex_norm)

        # --- Формування шляхів ---
        stem = Path(case_file).stem
        base_dir = Path(case_file).parent
        paths = {
            "edx": base_dir / f"{stem}_tasks_edx.txt",
            "json": base_dir / f"{stem}_tasks.json",
            "readable": base_dir / f"{stem}_tasks.docx",
        }

        # --- Збереження файлів ---
        with open(paths["json"], "w", encoding="utf-8") as f:
            json.dump({"exercises": exercises}, f, ensure_ascii=False, indent=2)

        self._write_edx_docx(paths["edx"], exercises)
        self._write_readable_docx(paths["readable"], exercises)

        if verbose:
            print("Saved:", str(paths["edx"]), str(paths["json"]), str(paths["readable"]))

        return str(paths["edx"]), str(paths["json"]), str(paths["readable"])
        
    # --------------------------------------------------------------------- #
    def _shuffle_and_relabel_options(self, ex: Dict[str, Any], rng: random.Random) -> Dict[str, Any]:
        opts = ex.get("options", []) or []
        rng.shuffle(opts)
        for i, o in enumerate(opts):
            o["letter"] = chr(ord("A") + i)
        ex["options"] = opts
        return ex

    def _load_case_studies(self, case_file: str) -> str:
        p = Path(case_file)
        if not p.exists():
            raise FileNotFoundError(f"Файл не знайдено: {case_file}")
        ext = p.suffix.lower()
        if ext not in (".json", ".docx"):
            raise ValueError(f"case_file має бути .json або .docx, отримано: {ext}")
        if ext == ".json":
            try:
                raw_data = json.loads(p.read_text(encoding="utf-8"))
            except json.JSONDecodeError as e:
                raise ValueError(f"Невалідний JSON у файлі {case_file}: {e}")

            # --- Збираємо ВЕСЬ текст із JSON ---
            def extract_all_text(data, path="", collected=None):
                if collected is None:
                    collected = []

                if isinstance(data, dict):
                    for k, v in data.items():
                        key_context = f"[{k}]" if path else k
                        # Додаємо назву ключа як контекст
                        if isinstance(v, str) and v.strip():
                            collected.append(f"--- {key_context} ---\n{v.strip()}")
                        elif isinstance(v, (list, dict)):
                            extract_all_text(v, path + key_context, collected)
                        # Ігноруємо bool, int, null
                elif isinstance(data, list):
                    for i, item in enumerate(data):
                        item_path = f"{path}[{i}]" if path else f"[{i}]"
                        if isinstance(item, str) and item.strip():
                            collected.append(f"--- {item_path} ---\n{item.strip()}")
                        elif isinstance(item, (dict, list)):
                            extract_all_text(item, item_path, collected)
                elif isinstance(data, str) and data.strip():
                    collected.append(f"--- {path or 'root'} ---\n{data.strip()}")

                return collected

            parts = extract_all_text(raw_data)
            if not parts:
                raise ValueError(f"JSON файл {case_file} не містить жодного текстового вмісту.")

            full_text = "\n\n".join(parts)
            # print(f"[DEBUG] Extracted {len(parts)} text block(s) from JSON. Total: {len(full_text)} chars")
            return full_text
        if ext == ".docx":
            if docx_reader is None:
                raise RuntimeError("Встановіть python-docx")
            d = docx_reader.Document(str(p))
            return "\n".join(para.text for para in d.paragraphs if para.text.strip())
        return ""

    # ------------------- КЛЮЧ: ВИКЛИК gpt-5-nano ------------------- #
    def _call_model_structured(
        self,
        case_text: str,
        n_questions: int,
        n_answers: int,
        n_correct: int,
        provided_questions: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        # --- SYSTEM PROMPT ---
        system_prompt = (
            "You are a seasoned professor and instructional designer specializing in university course and case-based learning. "
            "You design clear training exercises.\n"
            "CRITICAL RULES:\n"
            "• Use ONLY the content and themes provided in the case study text.\n"
            "• Do NOT add external facts, brands, or examples.\n"
            "• Stay fully within the context of the provided case file.\n"
            "• Paraphrase; do not copy sentences verbatim.\n"
            "• Return ONLY valid JSON. No explanations."
        )

        # --- USER PROMPT: ПОВНИЙ case_text ---
        # НЕ обрізаємо! gpt-5-nano приймає до ~32k токенів
        case_content = case_text.strip()
        if len(case_content) == 0:
            raise ValueError("case_file порожній або не прочитаний")

        if provided_questions:
            qlist = "\n".join(f"Q{i+1}: {q}" for i, q in enumerate(provided_questions))
            user_prompt = (
                f"CASE STUDY (your ONLY context — use this text only):\n"
                f"{case_content}\n\n"
                f"QUESTIONS TO ANSWER (use these EXACTLY as question text):\n"
                f"{qlist}\n\n"
                f"TASK:\n"
                f"- Generate EXACTLY {n_questions} multiple-choice exercises.\n"
                f"- Each must have EXACTLY {n_answers} options (A, B, C, ...).\n"
                f"- EXACTLY {n_correct} must be correct.\n"
                f"- For each exercise:\n"
                f"  • short_title: 3–7 word summary\n"
                f"  • question: COPY the original question EXACTLY\n"
                f"  • tip: 2-sentence hint with context from case\n"
                f"  • options: list with:\n"
                f"       - letter: 'A', 'B', etc.\n"
                f"       - text: ONLY the answer content — DO NOT start with 'A.', 'B.', etc.\n"
                f"       - is_correct: true/false\n"
                f"       - selected_fb: 2–3 sentences why correct\n"
                f"       - unselected_fb: 1–2 sentences why incorrect\n"
                f"- CRITICAL: In 'text' field — write ONLY the content. Example:\n"
                f"     \"text\": \"Data sharing violates policy\"\n"
                f"     NOT: \"A. Data sharing violates policy\"\n"
                f"- Output ONLY valid JSON. Start with {{\n"
            )
        else:
            user_prompt = (
                f"CASE STUDY (your ONLY context — use this text only):\n"
                f"{case_content}\n\n"
                f"TASK:\n"
                f"- Create EXACTLY {n_questions} multiple-choice exercises based ONLY on the case above.\n"
                f"- Each must have EXACTLY {n_answers} options.\n"
                f"- EXACTLY {n_correct} correct.\n"
                f"- For each exercise:\n"
                f"  • short_title: 3–7 word summary\n"
                f"  • question: full question in English\n"
                f"  • tip: 2-sentence hint with context\n"
                f"  • options: with:\n"
                f"       - letter: 'A', 'B', etc.\n"
                f"       - text: ONLY the answer content — DO NOT start with 'A.', 'B.', etc.\n"
                f"       - is_correct: true/false\n"
                f"       - selected_fb: 2–3 sentences why correct\n"
                f"       - unselected_fb: 1–2 sentences why incorrect\n"
                f"- CRITICAL: In 'text' field — write ONLY the content. Example:\n"
                f"     \"text\": \"Data sharing violates policy\"\n"
                f"     NOT: \"A. Data sharing violates policy\"\n"
                f"- Output ONLY valid JSON. Start with {{\n"
            )
        # --- ЛОГУВАННЯ (для перевірки) ---
        # print(f"\n[DEBUG] Case text length: {len(case_content)} chars")
        # print(f"[DEBUG] First 200 chars:\n{case_content[:200]}...\n")

        # print(system_prompt)
        # print("!!!!")
        # print(user_prompt)
        # --- Виклик моделі ---
        try:
            resp = self.client.responses.create(
                model=self.model,
                input=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
            )

            raw = getattr(resp, "output_text", None)
            if not raw:
                raw = resp.output[0].content[0].text if resp.output else ""

            raw = raw.strip()
            if not raw:
                raise ValueError("Empty response from model")

            # Витягуємо JSON
            start = raw.find("{")
            if start == -1:
                raise ValueError("No JSON start found")
            data = json.loads(raw[start:])

            if "exercises" not in data or not isinstance(data["exercises"], list):
                raise ValueError("Invalid structure: no exercises list")

            # --- Гарантія полів ---
            for ex in data["exercises"]:
                if not ex.get("short_title"):
                    q = ex.get("question", "")
                    ex["short_title"] = " ".join(q.split()[:6]) + ("..." if len(q.split()) > 6 else "")

                if not ex.get("tip") or len(ex["tip"].split()) < 8:
                    ex["tip"] = "Review the case facts. Pay attention to roles, timelines, and policies."

                for opt in ex.get("options", []):
                    if not opt.get("selected_fb"):
                        opt["selected_fb"] = (
                            "This is correct. It aligns with the evidence and rules in the case study."
                            if opt.get("is_correct") else
                            "This is incorrect. It contradicts a key fact or assumption in the case."
                        )
                    if not opt.get("unselected_fb"):
                        opt["unselected_fb"] = (
                            "You should select this — it is supported by the case."
                            if opt.get("is_correct") else
                            "Correctly skipped — this option is not justified."
                        )

            return data

        except Exception as e:
            print(f"LLM failed: {e}")
            print("Using fallback generator...")
            return self._generate_fallback_exercises(
                case_text=case_text,
                n_questions=n_questions,
                n_answers=n_answers,
                n_correct=n_correct,
                provided_questions=provided_questions
            )
            
    def _generate_fallback_exercises(
        self,
        case_text: str,
        n_questions: int,
        n_answers: int,
        n_correct: int,
        provided_questions: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        exercises = []
        # ← ВИПРАВЛЕНО: provided_questions, а не questions
        questions = provided_questions or [f"Question {i+1} from the case." for i in range(n_questions)]

        for i, q in enumerate(questions[:n_questions]):
            # short_title з питання
            words = q.split()[:6]
            short_title = " ".join(words) + ("..." if len(q.split()) > 6 else "")

            options = []
            correct_idx = random.sample(range(n_answers), n_correct)
            for j in range(n_answers):
                is_correct = j in correct_idx
                options.append({
                    "letter": chr(ord("A") + j),
                    "text": f"Option {chr(ord('A') + j)}{' (correct)' if is_correct else ''}",
                    "is_correct": is_correct,
                    "selected_fb": (
                        "This is correct because it directly follows from the case facts. "
                        "The evidence supports this conclusion without contradiction."
                    ) if is_correct else (
                        "This is incorrect. It contradicts the stated policy or timeline in the case. "
                        "Always double-check the source."
                    ),
                    "unselected_fb": (
                        "You should select this — it is the only option backed by evidence."
                    ) if is_correct else (
                        "Correctly skipped. This option introduces an assumption not present in the case."
                    ),
                })

            exercises.append({
                "short_title": short_title or f"Task {i+1}",
                "question": q,
                "tip": "Focus on the timeline and stakeholder roles. Look for explicit statements in the case.",
                "options": options
            })

        return {"exercises": exercises}    # --------------------------------------------------------------------- #
    @staticmethod
    def _validate_counts(n_answers: int, n_correct: int):
        if n_answers < 2:
            raise ValueError("n_answers >= 2")
        if not (1 <= n_correct < n_answers):
            raise ValueError("1 <= n_correct < n_answers")

    @staticmethod
    def _enforce_counts(ex: Dict[str, Any], n_answers: int, n_correct: int):
        opts = ex.get("options", []) or []
        opts = opts[:n_answers]
        for o in opts:
            o["is_correct"] = bool(o.get("is_correct", False))
        cur = sum(1 for o in opts if o["is_correct"])
        if cur > n_correct:
            for o in reversed(opts):
                if cur <= n_correct: break
                if o["is_correct"]:
                    o["is_correct"] = False
                    cur -= 1
        elif cur < n_correct:
            for o in opts:
                if cur >= n_correct: break
                if not o["is_correct"]:
                    o["is_correct"] = True
                    cur += 1
        while len(opts) < n_answers:
            opts.append({
                "letter": chr(ord("A") + len(opts)),
                "text": "Placeholder",
                "is_correct": False,
                "selected_fb": "Fallback.",
                "unselected_fb": "Added."
            })
        ex["options"] = opts[:n_answers]

    # --------------------------------------------------------------------- #
    def _write_edx_docx(self, path: Path, exercises: List[Dict[str, Any]]):
        txt_path = path.with_suffix(".txt")
        
        with open(txt_path, "w", encoding="utf-8") as f:
            for idx, ex in enumerate(exercises, start=1):
                f.write(f">>TASK {idx}: {ex['short_title']}||{ex['tip']}<<\n")
                f.write(f"Question: {ex['question']}\n\n")
                
                for o in ex["options"]:
                    marker = "[x]" if o["is_correct"] else "[ ]"
                    feedback_selected = o['selected_fb']
                    feedback_unselected = o['unselected_fb']
                    correctness = "Correct" if o["is_correct"] else "Incorrect"
                    
                    f.write(
                        f"{marker} {o['letter']}. {o['text']}\n"
                        f"    {{selected: {o['letter']} {correctness} – {feedback_selected}}}\n"
                        f"    {{unselected: {o['letter']} {'Incorrect to skip' if o['is_correct'] else 'Correct to skip'} – {feedback_unselected}}}\n"
                    )
                f.write("\n" + "-"*80 + "\n\n")
            # doc = Document()
        # for ex in exercises:
        #     doc.add_paragraph(f">>{ex['question']}||{ex['tip']}<<")
        #     for o in ex["options"]:
        #         marker = "[x]" if o["is_correct"] else "[ ]"
        #         doc.add_paragraph(f"{marker} {o['letter']}. {o['text']} {{selected: {o['letter']} {'Correct' if o['is_correct'] else 'Incorrect'} – {o['selected_fb']}}} {{unselected: {o['letter']} {'Incorrect to skip' if o['is_correct'] else 'Correct to skip'} – {o['unselected_fb']}}}")
        #     doc.add_paragraph("")
        # doc.save(path)

    def _write_readable_docx(self, path: Path, exercises: List[Dict[str, Any]]):
        doc = Document()
        doc.add_heading("Case Study Tasks", level=1)
        for i, ex in enumerate(exercises, start=1):
            doc.add_heading(f"Task {i}: {ex['short_title']}", level=2)
            p = doc.add_paragraph()
            p.add_run(ex["question"]).bold = True
            p.add_run(f" (Hint: {ex['tip']})").italic = True
            for o in ex["options"]:
                marker = "Correct" if o["is_correct"] else "Incorrect"
                doc.add_paragraph(f"  • {o['letter']}. {o['text']} — {marker}")
        doc.save(path)

    def _write_answers_json(self, path: Path, exercises: List[Dict[str, Any]]):
        data = []
        for i, ex in enumerate(exercises, start=1):
            correct = [o for o in ex["options"] if o["is_correct"]]
            data.append({
                "task": i,
                "question": ex["question"],
                "correct": [{"letter": o["letter"], "text": o["text"]} for o in correct]
            })
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def _write_answers_docx(self, path: Path, exercises: List[Dict[str, Any]]):
        doc = Document()
        doc.add_heading("Правильні відповіді", level=1)
        for i, ex in enumerate(exercises, start=1):
            doc.add_heading(f"Завдання {i}", level=2)
            doc.add_paragraph(ex["question"])
            for o in ex["options"]:
                if o["is_correct"]:
                    p = doc.add_paragraph()
                    p.add_run(f"{o['letter']}. {o['text']}").bold = True
        doc.save(path)

    @staticmethod
    def _coerce_bool(val) -> bool:
        return val in (True, "true", "True", 1, "1", "yes")

    @staticmethod
    def _first_nonempty(*vals) -> str:
        for v in vals:
            if isinstance(v, str) and v.strip():
                return v.strip()
        return ""

    def _normalize_exercise(self, ex: Dict[str, Any], n_answers: int) -> Dict[str, Any]:
        # Гарантуємо short_title
        short_title = self._first_nonempty(ex.get("short_title"), ex.get("title", ""))
        if not short_title:
            q = ex.get("question", "")
            words = q.split()[:6]
            short_title = " ".join(words) + ("..." if len(q.split()) > 6 else "")

        norm = {
            "short_title": short_title,
            "question": self._first_nonempty(ex.get("question", "")),
            "tip": self._first_nonempty(ex.get("tip", ""), "Read the case carefully."),
            "options": [],
        }
        raw_opts = ex.get("options", []) or []
        raw_opts = list(raw_opts)[:n_answers]
        letters = [chr(ord("A") + i) for i in range(n_answers)]
        for idx, raw in enumerate(raw_opts):
            raw = raw or {}
            letter = str(raw.get("letter") or letters[idx]).upper()[:1]
            text = self._first_nonempty(raw.get("text"), raw.get("answer"), f"Option {letter}")
            is_correct = self._coerce_bool(raw.get("is_correct", False))
            sel = self._first_nonempty(raw.get("selected_fb"), "Correct." if is_correct else "Incorrect.")
            unsel = self._first_nonempty(raw.get("unselected_fb"), "Select this." if is_correct else "Skip.")
            norm["options"].append({
                "letter": letter,
                "text": text,
                "is_correct": is_correct,
                "selected_fb": sel,
                "unselected_fb": unsel,
            })
        return norm