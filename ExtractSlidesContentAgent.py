#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LectureSlidesAgent — мінімалістичний агент класом.

Ініціалізація приймає шлях до .docx з розміткою:

  Slide 1
  Title: ...
  Content:
  - пункт 1
  - пункт 2
  Lecturer Notes:
  ...

Методи API (нічого не друкують і не створюють файлів):
- list_slides() -> list[dict]
- get_title(slide_index: int) -> str
- get_content(slide_index: int) -> list[str]
- get_notes(slide_index: int) -> str
- slides_count() -> int
- get_slide(slide_index: int) -> dict  # {"title", "content", "notes"}

Приклад використання:
    agent = LectureSlidesAgent("content.docx")
    all_slides = agent.list_slides()
    title_2 = agent.get_title(2)        # індексація з 1
    bullets_2 = agent.get_content(2)
    notes_2 = agent.get_notes(2)

Залежності: pip install python-docx
"""
from __future__ import annotations
import re
import json
from pathlib import Path
from typing import List, Dict, Any

from docx import Document  # pip install python-docx

# ------------------------------
# Регулярні вирази
# ------------------------------
_SLIDE_RE = re.compile(r"(?:^|\n)Slide\s+\d+\s*(.*?)(?=(?:\nSlide\s+\d+|\Z))", re.S)


def _read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # Дочитуємо таблиці (рядок із таб-розділювачем)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    # Нормалізація перенесень
    text = re.sub(r"\r\n?", "\n", text)
    return text


def _split_into_slides(text: str) -> List[str]:
    return _SLIDE_RE.findall(text)


def _extract_title(block: str) -> str:
    m = re.search(r"Title:\s*(.+)", block)
    return m.group(1).strip() if m else "(No Title)"


def _extract_content(block: str) -> List[str]:
    m = re.search(
        r"Content:\s*(.*?)(?=\n\s*(?:Lecturer Notes:|Visuals:|Slide\s+\d+|Title:|\Z))",
        block,
        flags=re.S,
    )
    if not m:
        return []
    raw = m.group(1).strip()
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    # Знімаємо стандартні маркери "- " або "• "
    bullets = [re.sub(r"^(?:[-\u2022])\s+", "", ln) for ln in lines]
    return bullets


def _extract_notes(block: str) -> str:
    pattern = re.compile(
        r"Lecturer Notes:\s*(.*?)(?=\n\s*Visuals:|\n\s*---|\n\s*Slide\s+\d+|\n\s*Title:|\Z)",
        flags=re.S,
    )
    parts = pattern.findall(block)
    text = "\n\n".join(p.strip() for p in parts if p.strip())
    text = re.sub(r"(?m)^- ", "", text)
    return text.strip()


class LectureSlidesAgent:
    """Агент, що парсить .docx при ініціалізації і дає методи доступу."""

    def __init__(self, path: str, text: str = None):
        self._path = Path(path)
        if not self._path.exists():
            raise FileNotFoundError(f"Файл не знайдено: {self._path}")

        # Визначаємо тип файлу за розширенням
        if self._path.suffix.lower() == ".json":
            # Завантаження з JSON
            with self._path.open("r", encoding="utf-8") as f:
                data = json.load(f)

            self._slides = [
                {
                    "title": slide["title"],
                    "content": slide["content"],
                    "notes": slide["notes"]
                }
                for slide in data.get("slides", [])
            ]
        else:
            # Існуюча логіка для .docx
            if text is None:
                text = _read_docx_text(str(self._path))
            blocks = _split_into_slides(text)
            self._slides = [
                {
                    "title": _extract_title(b),
                    "content": _extract_content(b),
                    "notes": _extract_notes(b),
                }
                for b in blocks
            ]

    # ------------------------------
    # Публічне API
    # ------------------------------
    def slides_count(self) -> int:
        return len(self._slides)

    def list_slides(self) -> List[Dict[str, Any]]:
        """Повертає список слайдів як dict із ключами: title, content, notes."""
        return list(self._slides)

    def _get(self, idx: int) -> Dict[str, Any]:
        if idx < 1 or idx > len(self._slides):
            raise IndexError(f"Невірний індекс слайду: {idx}. Доступно 1..{len(self._slides)}")
        return self._slides[idx - 1]

    def get_slide(self, slide_index: int) -> Dict[str, Any]:
        return dict(self._get(slide_index))

    def get_title(self, slide_index: int) -> str:
        return self._get(slide_index)["title"]

    def get_content(self, slide_index: int) -> List[str]:
        return list(self._get(slide_index)["content"])  # копія списку

    def get_notes(self, slide_index: int) -> str:
        return self._get(slide_index)["notes"]
        
    def save_to_json(self) -> None:
        """
        Зберігає всі слайди у JSON-файл.
    
        Автоматична назва: <source_docx>.json
        Наприклад:
            content.docx → content.json
            my_slides.docx → my_slides.json
    
        Файл зберігається в тій самій папці, що й оригінальний .docx.
        """
        # Автоматична назва: content.docx → content.json
        json_path = self._path.with_suffix(".json")
    
        data = {
            "source_docx": self._path.name,
            "total_slides": self.slides_count(),
            "slides": [
                {
                    "slide_number": i + 1,
                    "title": slide["title"],
                    "content": slide["content"],
                    "notes": slide["notes"]
                }
                for i, slide in enumerate(self._slides)
            ]
        }
    
        # Створюємо папки, якщо треба
        json_path.parent.mkdir(parents=True, exist_ok=True)
    
        with json_path.open("w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)


# Жодного CLI — тільки клас і методи.
