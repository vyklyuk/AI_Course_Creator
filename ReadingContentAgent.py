#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ReadingContentAgent — генерує текст лекції у форматі глави книги (для "reading" типу).

Новий метод: save_to_json() — зберігає структурований JSON поруч із .docx
"""

from typing import Optional, Tuple, List, Dict
from pathlib import Path
from docx import Document
import re
import json

# ------------------------- Difficulty Profiles -------------------------
DIFFICULTY_PROFILES = {
    "Beginner": (
        "- Audience: newcomers; avoid jargon or define on first use.\n"
        "- Scaffolding: strong. Use short sentences, step-by-step guidance.\n"
        "- Examples: add simple, concrete examples and analogies.\n"
        "- Explanations: define acronyms and terms inline.\n"
        "- Tone: friendly, clear, encouraging.\n"
    ),
    "Intermediate": (
        "- Audience: practitioners; moderate jargon with quick clarifications.\n"
        "- Scaffolding: medium. Focus on practical tips and workflows.\n"
        "- Examples: show realistic scenarios and trade-offs.\n"
        "- Explanations: assume basics known; clarify non-obvious parts.\n"
        "- Tone: professional, concise.\n"
    ),
    "Advanced": (
        "- Audience: experts; domain terminology without basic definitions.\n"
        "- Scaffolding: light. Emphasize trade-offs, assumptions, edge cases.\n"
        "- Examples: deep dives, performance and limitation notes.\n"
        "- Explanations: focus on rigor, references, comparatives.\n"
        "- Tone: precise, analytical.\n"
    ),
    "Executive": (
        "- Audience: non-technical decision-makers.\n"
        "- Focus: business outcomes, risk, cost, ROI, timelines.\n"
        "- Scaffolding: narrative summaries, KPIs, options & implications.\n"
        "- Explanations: minimal technical detail; clear decisions needed.\n"
        "- Tone: crisp, outcome-oriented.\n"
    ),
}

# ------------------------- Word count target -------------------------
def calc_target_words(minutes: int) -> int:
    return round(minutes * (8000 / 60.0))

# ------------------------- System prompt -------------------------
SYSTEM_TEMPLATE = """You are a professional textbook author writing university course.
Write in a clear, structured, educational style — like a chapter in a university textbook.

STYLE RULES:
1) Structure: Start with a short introduction (1–2 paragraphs), then 3–5 logical subsections with headings, end with a concise conclusion.
2) Headings: Use clear, titled subsections (e.g., "## 1. Predicate Crimes").
3) Tone: Educational, precise, engaging. Use active voice where appropriate.
4) Examples: Include 1–2 short real-world examples per major concept.
5) Clarity: Define terms on first use. Use short paragraphs (3–5 sentences).
6) No slides, no bullet lists under "Content:", no "Lecturer Notes:" — write in flowing prose.
7) Language: English only.
8) Word count: The entire chapter (excluding headings) must be within ±10% of the target.

OUTPUT FORMAT (plain text):
Chapter Title: [Lecture Title]

[Introduction...]

## 1. Subsection Title
[Body...]

## Key Takeaways
- Point 1
- Point 2

## Conclusion
[Wrap-up...]

---"""

# ------------------------- Build prompts -------------------------
def build_system(difficulty_level: str = "Beginner") -> str:
    profile = DIFFICULTY_PROFILES.get(difficulty_level, DIFFICULTY_PROFILES["Beginner"])
    return SYSTEM_TEMPLATE + f"\n\nDifficulty profile:\n{profile}"

def build_user_prompt(
    course_title: str,
    module_title: str,
    lecture_title: str,
    minutes: int,
    lecture_short_desc: str | None = None,
) -> str:
    target_words = calc_target_words(minutes)
    min_words = int(target_words * 0.9)
    max_words = int(target_words * 1.1)

    desc = (lecture_short_desc or "").strip()
    if desc:
        desc = f"Use this description as the foundation: {desc}"
    else:
        desc = "Focus on core concepts, practical implications, and learning outcomes."

    return f"""
Write a textbook chapter for:

Course: {course_title}
Module: {module_title}
Chapter: {lecture_title}

{desc}

WORD COUNT (strict):
- Total words in the chapter body (excluding headings and Key Takeaways bullets): {min_words}–{max_words} (target ≈ {target_words})
- Count only running prose text.

Ensure:
- 3–5 meaningful subsections
- At least one example per major concept
- Key Takeaways: 3–5 bullets
- Conclusion: ties back to real-world application
"""


# ------------------------- Text → Structured JSON -------------------------
def _parse_text_to_structure(text: str, lecture_title: str) -> Dict:
    """Парсить текст у структурований словник для JSON"""
    lines = [ln.rstrip() for ln in text.splitlines()]

    structure = {
        "chapter_title": lecture_title,
        "introduction": [],
        "sections": [],
        "key_takeaways": [],
        "conclusion": [],
        "metadata": {
            "source": "ReadingContentAgent",
            "generated_at": None,
            "word_count": 0
        }
    }

    current_section = None
    in_intro = True
    in_key_takeaways = False
    in_conclusion = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("Chapter Title:"):
            continue

        elif line.startswith("## "):
            title = line[3:].strip()
            if "Key Takeaways" in title:
                in_key_takeaways = True
                in_intro = in_conclusion = False
                continue
            elif "Conclusion" in title:
                in_conclusion = True
                in_key_takeaways = in_intro = False
                continue
            else:
                if in_intro:
                    in_intro = False
                if current_section:
                    structure["sections"].append(current_section)
                current_section = {"title": title, "paragraphs": []}
                in_key_takeaways = in_conclusion = False
                continue

        elif line.startswith(("- ", "• ")):
            if in_key_takeaways:
                structure["key_takeaways"].append(line[2:].strip())
            # ігноруємо в інших місцях
            continue

        else:
            if in_intro:
                structure["introduction"].append(line)
            elif in_conclusion:
                structure["conclusion"].append(line)
            elif current_section:
                current_section["paragraphs"].append(line)

    if current_section:
        structure["sections"].append(current_section)

    # Підрахунок слів (тільки проза)
    prose = " ".join(
        structure["introduction"] +
        [p for sec in structure["sections"] for p in sec["paragraphs"]] +
        structure["conclusion"]
    )
    structure["metadata"]["word_count"] = len(prose.split())
    structure["metadata"]["generated_at"] = __import__("datetime").datetime.now().isoformat()

    return structure


# ------------------------- Agent -------------------------
class ReadingContentAgent:
    def __init__(self, client, model: str = "gpt-5-nano"):
        self.client = client
        self.model = model
        self.last_status = ""
        self.last_docx_path = None
        self.last_text = None

    def generate(
        self,
        course_title: str,
        module_title: str,
        lecture_title: str,
        lecture_short_desc: str | None = None,
        difficulty_level: str = "Beginner",
        minutes: int = 10,
        save_path: Optional[str] = None,
        verbose: bool = True,
    ) -> Tuple[str, str]:
        MAX_RETRIES = 5
        TOL = 0.10

        def _count_body_words(text: str) -> int:
            cleaned = re.sub(r"^Chapter Title:.*$", "", text, flags=re.M)
            cleaned = re.sub(r"^##.*$", "", cleaned, flags=re.M)
            cleaned = re.sub(r"^(?:-|\u2022)\s.*$", "", cleaned, flags=re.M)
            cleaned = re.sub(r"\n\s*\n", "\n", cleaned)
            words = cleaned.split()
            return len([w for w in words if w.strip()])

        target_words = calc_target_words(minutes)
        min_words = int(target_words * (1 - TOL))
        max_words = int(target_words * (1 + TOL))

        print(lecture_title, "!!!!!!!!!!!")
        system = build_system(difficulty_level)
        base_user = build_user_prompt(course_title, module_title, lecture_title, minutes, lecture_short_desc)

        if verbose:
            print(f"Generating reading content for: {lecture_title}")
            print(f"Target words: ~{target_words} (±10% → {min_words}–{max_words})")

        text = ""
        success = False
        current_words = 0

        for attempt in range(1, MAX_RETRIES + 1):
            user_prompt = base_user
            if attempt > 1:
                # Тільки при повторі — показуємо "Attempt N"
                print(f"Attempt {attempt}: adjusting word count...")
                if current_words < min_words:
                    delta = min_words - current_words
                    user_prompt += f"\n\nADJUST: EXPAND by ~{delta} words. Add depth, examples."
                else:
                    delta = current_words - max_words
                    user_prompt += f"\n\nADJUST: CONDENSE by ~{delta} words. Tighten prose."

            try:
                # print(system)
                # print(user_prompt)
                response = self.client.responses.create(
                    model=self.model,
                    input=[
                        {"role": "system", "content": system},
                        {"role": "user", "content": user_prompt},
                    ],
                )
                text = getattr(response, "output_text", None) or response.output[0].content[0].text
            except Exception as e:
                text = f"[ERROR: {e}]"
                if verbose:
                    print(f"API error: {e}")
                break

            current_words = _count_body_words(text)
            if min_words <= current_words <= max_words:
                success = True
                self.last_status = f"OK: {current_words} words (target {min_words}–{max_words})"
                if verbose:
                    print(f"Success: Generated: {current_words} words")
                break
            else:
                if verbose:
                    print(f"Attempt {attempt}: {current_words} words → retrying...")

        if not success and verbose:
            print(f"Warning: FAILED after {MAX_RETRIES} attempts: {current_words} words")

        # === ЗБЕРЕЖЕННЯ DOCX ===
        if not save_path:
            safe_title = re.sub(r"[^\w\s-]", "", lecture_title).strip()
            safe_title = re.sub(r"\s+", "_", safe_title)
            save_path = f"{safe_title}_reading.docx"

        save_path = Path(save_path)
        save_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        for line in text.splitlines():
            if line.strip():
                if line.startswith("Chapter Title:"):
                    doc.add_heading(line.replace("Chapter Title:", "").strip(), level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith(("Key Takeaways", "Conclusion")):
                    doc.add_heading(line.strip(), level=2)
                elif line.strip().startswith(("- ", "• ")):
                    doc.add_paragraph(line.strip(), style="List Bullet")
                else:
                    doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph()
        doc.save(str(save_path))

        self.last_text = text
        self.last_docx_path = str(save_path)
        # print(text)

        return text, str(save_path)
        
    # ------------------------- НОВИЙ МЕТОД: JSON -------------------------
    def save_to_json(self) -> str:
        """
        Зберігає структурований контент у JSON.
        Назва: <docx_name>.json (наприклад, predicate_reading.docx → predicate_reading.json)
        Повертає шлях до JSON.
        """
        if not self.last_text or not self.last_docx_path:
            raise RuntimeError("Спочатку виклич generate()")

        # Автоматична назва
        docx_path = Path(self.last_docx_path)
        json_path = docx_path.with_suffix(".json")

        # Парсимо текст у структуру
        lecture_title = re.search(r"Chapter Title:\s*(.+)", self.last_text)
        lecture_title = lecture_title.group(1).strip() if lecture_title else "Untitled Chapter"

        structure = _parse_text_to_structure(self.last_text, lecture_title)

        # Зберігаємо
        json_path.parent.mkdir(parents=True, exist_ok=True)
        with json_path.open("w", encoding="utf-8") as f:
            json.dump(structure, f, ensure_ascii=False, indent=2)

        return str(json_path)


# ------------------------- Приклад використання -------------------------
if __name__ == "__main__":
    from openai import OpenAI

    agent = ReadingContentAgent(client, model="gpt-5-nano")

    text, docx_path = agent.generate(
        course_title="AML Uncovered",
        module_title="Module 1 – AML Foundations",
        lecture_title="Predicate Crimes and Money Laundering",
        lecture_short_desc="Explain how criminal proceeds are generated...",
        difficulty_level="Intermediate",
        minutes=8,
        save_path="output/readings/predicate_reading.docx"
    )
    print(f"DOCX: {docx_path}")

    json_path = agent.save_to_json()
    print(f"JSON: {json_path}")