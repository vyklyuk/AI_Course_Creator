import json
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class CurriculaCourseAgent:
    def __init__(self, client, model):
        self.client = client
        self.model = model

    def _validate_structure(self, structure: Dict, total_minutes: int, exam_minutes: int, modules_count: int, exam_format: dict) -> bool:
        """Validates the course structure for correctness"""
        try:
            # 1. Basic fields
            if not structure.get('modules') or len(structure['modules']) != modules_count:
                print(f"ERROR: Wrong number of modules: {len(structure.get('modules', []))} ≠ {modules_count}")
                return False

            if structure.get('total_minutes', 0) != total_minutes:
                print(f"ERROR: total_minutes = {structure.get('total_minutes')} ≠ {total_minutes}")
                return False

            if structure.get('exam_minutes', 0) != exam_minutes:
                print(f"ERROR: exam_minutes = {structure.get('exam_minutes')} ≠ {exam_minutes}")
                return False

            # 2. Module duration validation
            calculated_module_time = 0
            num_case_questions = exam_format.get("num_case_questions", 1)
            min_reading_minutes = 20
            max_reading_minutes = num_case_questions * 10

            for module in structure['modules']:
                module_time = 0

                # Videos: 6–10 min
                for v in module.get('videos', []):
                    d = v.get('duration_minutes', 0)
                    if not (6 <= d <= 10):
                        print(f"ERROR: Video duration {d} min (must be 6–10)")
                        return False
                    module_time += d

                # Readings: 20–max_reading_minutes
                for r in module.get('readings', []):
                    d = r.get('duration_minutes', 0)
                    if not (min_reading_minutes <= d <= max_reading_minutes):
                        print(f"ERROR: Reading {d} min (must be {min_reading_minutes}–{max_reading_minutes})")
                        return False
                    module_time += d

                # Case Studies: duration = questions * 10
                for c in module.get('case_studies', []):
                    expected = len(c.get('questions', [])) * 10
                    actual = c.get('duration_minutes', 0)
                    if actual != expected:
                        print(f"ERROR: Case study: {actual} ≠ {expected} (expected {len(c.get('questions', []))} × 10)")
                        return False
                    module_time += actual

                # Quiz: 2 min per question
                for q in module.get('quiz', []):
                    expected = 2
                    actual = q.get('duration_minutes', 0)
                    if actual != expected:
                        print(f"ERROR: Quiz question: {actual} ≠ {expected}")
                        return False
                    module_time += actual

                if module_time == 0:
                    print("ERROR: Module has no duration")
                    return False
                calculated_module_time += module_time

            # 3. Exam
            exam_time = 0
            case_studies = structure['exam'].get('case_studies', [])
            if len(case_studies) > 1:
                print(f"ERROR: Exam has {len(case_studies)} case studies. Must be 0 or 1.")
                return False
            
            for c in case_studies:
                num_questions = len(c.get('questions', []))
                if num_questions != exam_format.get("num_case_questions", 0):
                    print(f"ERROR: Exam case study has {num_questions} questions, expected {exam_format.get('num_case_questions', 0)}")
                    return False
                expected = num_questions * 10
                actual = c.get('duration_minutes', 0)
                if actual != expected:
                    print(f"ERROR: Exam case study duration {actual} ≠ {expected}")
                    return False
                exam_time += actual

            for q in structure['exam'].get('quiz', []):
                expected = 2
                actual = q.get('duration_minutes', 0)
                if actual != expected:
                    print(f"ERROR: Exam question: {actual} ≠ {expected}")
                    return False
                exam_time += actual

            total_calculated = calculated_module_time + exam_time
            if abs(total_calculated - total_minutes) > 120:
                print(f"ERROR: Total time {total_calculated} ≠ {total_minutes} (allow ±120 min)")
                return False

            print("Validation passed successfully!")
            return True

        except Exception as e:
            print(f"Validation error: {e}")
            return False
    def _hours_to_minutes(self, hours: float) -> int:
        return int(hours * 60)

    def _calculate_quiz_questions(self, duration_minutes: int) -> int:
        return duration_minutes // 2

    def _extract_docx_text(self, docx_file_path: str) -> str:
        """Extract text from a .docx file."""
        try:
            doc = Document(docx_file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            print(f"Word loadad: {docx_file_path}")
            return "\n".join(full_text)
        except Exception as e:
            print(f"ERROR: Failed to read .docx file: {e}")
            return ""
   

   
    def generate_course_structure(self, 
                                 course_name: str,
                                 level: str,
                                 total_hours: float,
                                 course_description: str,
                                 modules_count: int,
                                 videos_per_module: int,
                                 readings_per_module: int,
                                 case_studies: bool,
                                 num_case_questions: int,
                                 quizzes: bool,
                                 exam_format: dict,
                                 learning_objectives: Optional[Dict[str, int]] = None,
                                 docx_file_path: Optional[str] = None,
                                 max_retries: int = 3) -> dict:
         # Форматування рівнів зі зірочками
        def _format_level(level: int) -> str:
            return "★" * level + "☆" * (5 - level)
            
        if learning_objectives:
            objectives_lines = []
            for obj, lvl in learning_objectives.items():
                if not (1 <= lvl <= 5):
                    raise ValueError(f"Learning objective level must be 1–5, got {lvl} for {obj}")
                stars = _format_level(lvl)
                objectives_lines.append(f"{obj}\t{lvl}\t{stars}")
            learning_objectives_str = "\n".join(objectives_lines)
        else:
            learning_objectives_str = "None provided"
        
        # Calculate total minutes and debug
        total_minutes = self._hours_to_minutes(total_hours)
        # print(f"DEBUG: Input total_hours={total_hours}, calculated total_minutes={total_minutes}")
        if total_minutes == 0:
            print(f"WARNING: total_minutes is 0, using fallback calculation: {int(total_hours * 60)}")
            total_minutes = int(total_hours * 60)
        
        # Calculate exam time based on exam_format
        exam_minutes = (exam_format.get("num_case_questions", 0) * 10) + (exam_format.get("num_quiz_questions", 0) * 2)
        minutes_per_module = (total_minutes - exam_minutes) // modules_count

        min_reading_minutes = 20
        max_reading_minutes = exam_format.get("num_case_questions", 5) * 10  # 10 хв на питання
        if max_reading_minutes < min_reading_minutes:
            max_reading_minutes = min_reading_minutes  # Захист
            
        # Calculate quiz questions per module (rule #10)
        quiz_questions_per_module = self._calculate_quiz_questions(15) if quizzes else 0

        # Extract text from docx if provided
        docx_content = self._extract_docx_text(docx_file_path) if docx_file_path else ""

        system_prompt = """
        You are an EXPERT curriculum designer. Your task is to create a PRECISE course structure and return VALID JSON ONLY. Do not include any text, comments, markdown, backticks, or explanations outside the JSON object. Follow the provided course details, CRITICAL RULES, and the user-provided course structure description (from a Word document) strictly. Ensure that 'questions' in case_studies and exam.case_studies are lists of strings, not dictionaries. Ensure 'total_minutes' is exactly total_hours * 60.
        The user-provided course structure description (from a Word document) is the PRIMARY SOURCE for the course structure. Analyze the description to extract and prioritize:
        - Module titles and their order.
        - Topics or titles for videos, readings, case studies, and quizzes.
        - Learning objectives, themes, or specific content mentioned.
        - Any indicated structure (e.g., number of modules, types of activities).
        Structure the course to closely match the provided description, adapting it to fit the JSON format and CRITICAL RULES. If the description lacks details (e.g., specific titles, descriptions, or questions), generate appropriate content based on the course_name, level, course_description, and context. If no description is provided, use the course_description and parameters to create a complete structure.

        Enhance the provided course_description to make it more engaging, professional, and structured. Expand it to 6-8 sentences, adding 2-3 sentences to highlight the course's value and target audience. Emphasize practical applications and align with the course name, level, and objectives. If the Word document provides a course description, incorporate its key points into the enhanced description.

        For each case study in modules and exam, include a 'duration_minutes' field calculated as the number of questions * 10. For each quiz in modules and exam, include a 'duration_minutes' field calculated as the number of questions * 2.

LEARNING OBJECTIVES LEVELS (Bloom's Taxonomy):
1 = Remember (recall facts)
2 = Understand (explain concepts)
3 = Apply (use in new situations)
4 = Analyze (break down, compare)
5 = Evaluate/Create (judge, design)

Each learning objective has a target level (1–5). You MUST tailor ALL content to match this level:
- Video/reading descriptions: use vocabulary and depth matching the level
- Case studies: complexity of scenario and questions must match the level
- Quiz/exam questions: cognitive demand must align with the level
- Examples: simple recall (1) vs. real-world analysis (4) vs. policy design (5)


        Output format:
        {
            "course_name": "<course_name>",
            "level": "<level>",
            "total_hours": <total_hours>,
            "total_minutes": <total_minutes>,
            "exam_minutes": <exam_minutes>,
            "description": "<enhanced_course_description>",
            "learning_objectives": [
                {"name": "Objective Name", "level": 4, "level_stars": "★★★★☆"},
                ...
            ],
            "modules": [
                {
                    "module_number": 1,
                    "title": "Module Title",
                    "videos": [
                        {"title": "Video Title", "duration_minutes": 6, "description": "Detailed 3-5 sentence description with examples and processes"},
                        ...
                    ],
                    "readings": [
                        {"title": "Reading Title", "duration_minutes": 30, "description": "Detailed 6-7 sentence description"},
                        ...
                    ],
                    // duration_minutes must be >= 20 and <= (num_case_questions * 10)
                    "case_studies": [
                        {"title": "Case Study Title", "description": "Detailed 6-7 sentence description", "learning_outcomes": ["Outcome 1", ...], "questions": ["Question 1", ...], "duration_minutes": <num_questions * 10>},
                        ...
                    ],
                    "quiz": [
                        {"question": "Quiz Question", "options": ["A", "B", "C", "D"], "answer": "A", "duration_minutes": <num_questions * 2>},
                        ...
                    ]
                },
                ...
            ],
            "exam": {
                "case_studies": [
                    {"title": "Exam Case Study", "description": "Detailed 6-7 sentence description", "learning_outcomes": ["Outcome 1", ...], "questions": ["Question 1", ...], "duration_minutes": <num_questions * 10>},
                    ...
                ],
                "quiz": [
                    {"question": "Exam Quiz Question", "options": ["A", "B", "C", "D"], "answer": "A", "duration_minutes": <num_questions * 2>},
                    ...
                ],
                "passing_threshold": "<passing_threshold>"
            }
        }

        Return ONLY the JSON object. Do NOT include any additional text, markdown, or backticks.
        """

        exam_format_str = ""
        if exam_format.get("case_studies"):
            exam_format_str += f"1 Case Study with {exam_format.get('num_case_questions', 0)} questions"
        if exam_format.get("quizzes"):
            if exam_format_str:
                exam_format_str += " | "
            exam_format_str += f"{exam_format.get('num_quiz_questions', 0)} quiz questions"
        exam_format_str += f", passing threshold {exam_format.get('passing_threshold', '70%')}"


        user_prompt = f"""
        Course Name: {course_name}
        Level: {level}
        Total Duration: {total_hours} hours ({total_minutes} minutes)
        Description: {course_description}
        Modules: {modules_count}, each with {videos_per_module} videos and {readings_per_module} readings
        Readings: {readings_per_module} per module, duration {min_reading_minutes}–{max_reading_minutes} minutes each
        Case Studies: {'Yes' if case_studies else 'No'}, {num_case_questions} questions per case
        Quizzes: {'Yes' if quizzes else 'No'}, {quiz_questions_per_module} questions per quiz
        Exam Format: {exam_format_str}
        Exam Duration: {exam_minutes} minutes

        LEARNING OBJECTIVES (Name → Level → Bloom's Stars):
        {learning_objectives_str}

        User-Provided Course Structure Description (from Word document): 
        {docx_content if docx_content else 'None provided'}

        CRITICAL RULES:
        1. Minutes/Module: EXACTLY {minutes_per_module}
        2. Total = EXACTLY {total_minutes} MINUTES
        3. EACH MODULE MUST HAVE EXACTLY {videos_per_module} VIDEOS
        4. EVERY VIDEO: 6-10 MINUTES MAX
        5. MIX Videos + Readings (60% video, 40% reading)
        6. VIDEO DESCRIPTION: 3-5 sentences
        7. READING DESCRIPTION: 6-7 sentences
        8. CASE STUDY DESCRIPTION: 6-7 sentences + LEARNING OUTCOMES list
        9. EACH CASE STUDY: EXACTLY {num_case_questions} QUESTIONS
        10. EACH QUIZ: EXACTLY {quiz_questions_per_module} QUESTIONS
        11. EXAM RULE:
                - The exam must contain AT MOST 1 case study.
                - If case studies are included, there must be EXACTLY 1 case study.
                - This single case study must have EXACTLY {num_case_questions} questions.
                - Each question = 10 minutes → duration_minutes = {num_case_questions} * 10
        12. ALL DURATIONS IN MINUTES (integers only)
        13. ALL TEXT IN ENGLISH ONLY
        14. DISTRIBUTE LEARNING OBJECTIVES ACROSS MODULES: Each objective must appear in at least one module. 
        15: Content complexity MUST match the level:
            - Level 1: Define, list, recall
            - Level 2: Explain, summarize
            - Level 3: Apply rules to examples
            - Level 4: Analyze transactions, detect patterns
            - Level 5: Design policies, evaluate systems
        16. READINGS DURATION: 
            - MINIMUM 20 minutes per reading
            - MAXIMUM = duration of one case study (i.e., {num_case_questions} × 10 minutes)
            - Use integer values only
            - Example: if num_case_questions = 5 → max 50 min → reading duration = 20–50 min
        """
        
        
        # print(system_prompt)
        # print("*"*100)
        # print(user_prompt)
        for attempt in range(1, max_retries + 1):
            print(f"\nGENERATING COURSE — attempt {attempt}/{max_retries}")

            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ]
                )
                raw_content = response.choices[0].message.content
                structure = json.loads(raw_content)

                # Fix required fields
                structure['total_minutes'] = total_minutes
                structure['exam_minutes'] = exam_minutes
                if learning_objectives:
                    structure["learning_objectives"] = [
                        {"name": n, "level": l, "level_stars": _format_level(l)}
                        for n, l in learning_objectives.items()
                    ]
                else:
                    structure["learning_objectives"] = []

                # VALIDATION
                if self._validate_structure(structure, total_minutes, exam_minutes, modules_count, exam_format):
                    print(f"SUCCESS on attempt {attempt}!")
                    return structure
                else:
                    print(f"Attempt {attempt} failed. Retrying...")

            except Exception as e:
                print(f"Error on attempt {attempt}: {e}")

        # --- END OF LOOP ---
        print("All attempts failed. Returning fallback.")
        return {
            "course_name": course_name,
            "level": level,
            "total_hours": total_hours,
            "total_minutes": total_minutes,
            "exam_minutes": exam_minutes,
            "description": course_description,
            "learning_objectives": [],
            "modules": [],
            "exam": {
                "case_studies": [],
                "quiz": [],
                "passing_threshold": exam_format.get("passing_threshold", "70%")
            }
        }    
        # try:
        #     # print(system_prompt)
        #     # print(user_prompt)
        #     response = self.client.chat.completions.create(
        #         model=self.model,
        #         messages=[
        #             {"role": "system", "content": system_prompt},
        #             {"role": "user", "content": user_prompt}
        #         ]
        #     )
        #     raw_content = response.choices[0].message.content
        #     # Log raw response for debugging
        #     structure = json.loads(raw_content)
        #     # Ensure total_minutes is correct in the returned structure

        #     structure['total_minutes'] = total_minutes
        #     structure['exam_minutes'] = exam_minutes
            
        #     # ДОДАЄМО level_stars для кожного об’єктиву
        #     if learning_objectives:
        #         structure["learning_objectives"] = [
        #             {
        #                 "name": name,
        #                 "level": level,
        #                 "level_stars": "★" * level + "☆" * (5 - level)
        #             }
        #             for name, level in learning_objectives.items()
        #         ]
        #     else:
        #         structure["learning_objectives"] = []   
        #     return structure
        # except json.JSONDecodeError as e:
        #     print(f"ERROR: Failed to parse API response as JSON: {e}")
        #     print(f"Raw response: {raw_content}")
        #     # Fallback: Return a minimal valid structure
        #     return {
        #         "course_name": course_name,
        #         "level": level,
        #         "total_hours": total_hours,
        #         "total_minutes": total_minutes,  # Use calculated value
        #         "exam_minutes": exam_minutes,
        #         "description": course_description,  # Use original description as fallback
        #         "modules": [],
        #         "exam": {
        #             "case_studies": [],
        #             "quiz": [],
        #             "passing_threshold": exam_format.get("passing_threshold", "70%")
        #         }
        #     }
        # except Exception as e:
        #     print(f"ERROR: Unexpected error in API call: {e}")
        #     # Fallback: Return a minimal valid structure
        #     return {
        #         "course_name": course_name,
        #         "level": level,
        #         "total_hours": total_hours,
        #         "total_minutes": total_minutes,  # Use calculated value
        #         "exam_minutes": exam_minutes,
        #         "description": course_description,  # Use original description as fallback
        #         "modules": [],
        #         "exam": {
        #             "case_studies": [],
        #             "quiz": [],
        #             "passing_threshold": exam_format.get("passing_threshold", "70%")
        #         }
        #     }

    def print_professional_outline(self, structure: Dict):
        """Enhanced output with ALL fixes, including duration for each learning element and total time summary"""
        total_minutes = structure.get('total_minutes', 0)
        exam_minutes = structure.get('exam_minutes', 0)
    
        print("\n" + "═"*120)
        print(f"COURSE {structure['course_name']} | {structure['level']} | {total_minutes} MINUTES ({structure['total_hours']} hours)")
        print(f"Description: {structure.get('description', 'N/A')}")
    
        # Learning Objectives
        print(f"LEARNING OBJECTIVES (Bloom's Taxonomy Level):")
        for obj in structure.get('learning_objectives', []):
            print(f"   • {obj['name']} → Level {obj['level']} {obj['level_stars']}")
    
        print(f"Exam Duration: {exam_minutes} minutes")
        print("═"*120)
    
        total_course_minutes = 0  # Track total time for all learning elements
    
        for module in structure['modules']:
            module_total_minutes = 0
            print(f"\nModule {module['module_number']}: {module['title']}")
    
            # Videos
            print(f"  Videos ({len(module['videos'])}):")
            for video in module['videos']:
                video_duration = video.get('duration_minutes', 0)
                module_total_minutes += video_duration
                print(f"    - {video['title']} ({video_duration} min): {video.get('description', 'N/A')}")
    
            # Readings
            print(f"  Readings ({len(module['readings'])}):")
            for reading in module['readings']:
                reading_duration = reading.get('duration_minutes', 0)
                module_total_minutes += reading_duration
                print(f"    - {reading['title']} ({reading_duration} min): {reading.get('description', 'N/A')}")
    
            # Case Studies
            if module.get('case_studies'):
                print(f"  Case Studies ({len(module['case_studies'])}):")
                for case in module['case_studies']:
                    case_duration = case.get('duration_minutes', len(case.get('questions', [])) * 10)
                    module_total_minutes += case_duration
                    print(f"    - {case['title']} ({case_duration} min): {case.get('description', 'N/A')}")
                    print(f"      Learning Outcomes: {', '.join(case.get('learning_outcomes', ['N/A']))}")
                    questions = case.get('questions', [])
                    question_texts = [q['text'] if isinstance(q, dict) and 'text' in q else str(q) for q in questions]
                    print(f"      Questions ({len(question_texts)}): {', '.join(question_texts) if question_texts else 'N/A'}")
    
            # Quizzes
            if module.get('quiz'):
                print(f"  Quiz ({len(module['quiz'])} questions):")
                quiz_duration = sum(q.get('duration_minutes', 2) for q in module['quiz'])
                module_total_minutes += quiz_duration
                for quiz in module['quiz']:
                    quiz_question_duration = quiz.get('duration_minutes', 2)
                    print(f"    - {quiz['question']} ({quiz_question_duration} min)")
                print(f"    Quiz Total Duration: {quiz_duration} min")
    
            # Module time summary
            print(f"  Module Total Duration: {module_total_minutes} min")
            total_course_minutes += module_total_minutes
    
        # Final Exam
        print("\n" + "═"*120)
        print("FINAL EXAM:")
        exam_total_minutes = 0
        if structure['exam'].get('case_studies'):
            print(f"  Case Studies ({len(structure['exam']['case_studies'])}):")
            for case in structure['exam']['case_studies']:
                case_duration = case.get('duration_minutes', len(case.get('questions', [])) * 10)
                exam_total_minutes += case_duration
                print(f"    - {case['title']} ({case_duration} min): {case.get('description', 'N/A')}")
                print(f"      Learning Outcomes: {', '.join(case.get('learning_outcomes', ['N/A']))}")
                questions = case.get('questions', [])
                question_texts = [q['text'] if isinstance(q, dict) and 'text' in q else str(q) for q in questions]
                print(f"      Questions ({len(question_texts)}): {', '.join(question_texts) if question_texts else 'N/A'}")
    
        if structure['exam'].get('quiz'):
            print(f"  Quiz ({len(structure['exam']['quiz'])} questions):")
            quiz_duration = sum(q.get('duration_minutes', 2) for q in structure['exam']['quiz'])
            exam_total_minutes += quiz_duration
            for quiz in structure['exam']['quiz']:
                quiz_question_duration = quiz.get('duration_minutes', 2)
                print(f"    - {quiz['question']} ({quiz_question_duration} min)")
            print(f"    Quiz Total Duration: {quiz_duration} min")
    
        print(f"  Passing Threshold: {structure['exam']['passing_threshold']}")
        print(f"  Exam Total Duration: {exam_total_minutes} min")
    
        # Course time summary
        total_course_minutes += exam_total_minutes
        print("\n" + "═"*120)
        print(f"Total Course Duration (Modules + Exam): {total_course_minutes} min")
        print(f"Expected Total Duration (from input): {total_minutes} min")
        print("═"*120)


    # Word Creation    
    def create_course_docx(self, json_file_path: str, output_docx: str = "COURSE.docx"):
        """
        🎓 JSON → BEAUTIFUL DOCX COURSE BUILDER
        Creates a Word document from a JSON course structure, including duration for each learning element and total time summary.
        Ensures output file has .docx extension.
        """
        # Ensure output file has .docx extension
        if not output_docx.lower().endswith('.docx'):
            output_docx = f"{output_docx}.docx"
        
        # Load JSON
        try:
            with open(json_file_path, 'r', encoding='utf-8') as f:
                structure = json.load(f)
        except Exception as e:
            print(f"ERROR: Failed to load JSON file: {e}")
            return None
        
        # Create document
        doc = Document()
        
        # 🎨 SETUP STYLES
        self.setup_styles(doc)
        
        # 📄 TITLE PAGE
        self.add_title_page(doc, structure)
        doc.add_page_break()
        
        # 📋 TABLE OF CONTENTS
        self.add_table_of_contents(doc, structure)
        doc.add_page_break()
        
        # 🧩 MODULES
        self.add_modules(doc, structure)
        
        # 📜 FINAL EXAM
        self.add_final_exam(doc, structure)
        
        # Save
        try:
            doc.save(output_docx)
            print(f"✅ DOCX CREATED: {output_docx}")
            return output_docx
        except Exception as e:
            print(f"ERROR: Failed to save DOCX: {e}")
            return None

    def setup_styles(self, doc):
        """Define beautiful styles for the document"""
        # Heading 1 (Module)
        h1 = doc.styles.add_style('ModuleTitle', WD_STYLE_TYPE.PARAGRAPH)
        h1.font.name = 'Arial'
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 102, 204)  # Blue
        
        # Heading 2 (Element)
        h2 = doc.styles.add_style('ElementTitle', WD_STYLE_TYPE.PARAGRAPH)
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        # Description
        desc = doc.styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
        desc.font.name = 'Times New Roman'
        desc.font.size = Pt(11)

    def add_title_page(self, doc, structure):
        """Add title page with course details"""
        title = doc.add_heading(structure.get('course_name', 'Untitled Course'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(32)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
        subtitle = doc.add_paragraph(structure.get('description', 'No description available'))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].italic = True

        if structure.get('learning_objectives'):
            doc.add_paragraph("LEARNING OBJECTIVES (Bloom's Taxonomy):", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = doc.add_table(rows=len(structure['learning_objectives']), cols=2)
            table.style = 'Table Grid'
            for i, obj in enumerate(structure['learning_objectives']):
                row = table.rows[i].cells
                row[0].text = obj['name']
                level_para = row[1].add_paragraph()
                # level_run = level_para.add_run(f"Level {obj['level']} ")
                # level_run.bold = True
                stars_run = level_para.add_run(obj['level_stars'])
                stars_run.font.size = Pt(14)
                # Колір за рівнем
                color = {
                    1: RGBColor(100, 100, 100),
                    2: RGBColor(0, 100, 200),
                    3: RGBColor(0, 150, 0),
                    4: RGBColor(255, 165, 0),
                    5: RGBColor(200, 0, 0)
                }.get(obj['level'], RGBColor(0, 0, 0))
                stars_run.font.color.rgb = color
        
        info = doc.add_paragraph(f"Level: {structure.get('level', 'N/A')} | Total: {structure.get('total_minutes', 0)} min ({structure.get('total_hours', 0)} hours)")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.runs[0].font.size = Pt(12)

    def add_table_of_contents(self, doc, structure):
        """Add table of contents with module and exam durations, plus total course duration"""
        toc = doc.add_heading('TABLE OF CONTENTS', level=1)
        toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Calculate total course duration
        total_duration = 0
        exam_duration = 0
        for module in structure.get('modules', []):
            module_duration = sum(v.get('duration_minutes', 0) for v in module.get('videos', []))
            module_duration += sum(r.get('duration_minutes', 0) for r in module.get('readings', []))
            module_duration += sum(c.get('duration_minutes', 0) for c in module.get('case_studies', []))
            module_duration += sum(q.get('duration_minutes', 0) for q in module.get('quiz', []))
            total_duration += module_duration
        exam = structure.get('exam', {})
        exam_duration += sum(c.get('duration_minutes', 0) for c in exam.get('case_studies', []))
        exam_duration += sum(q.get('duration_minutes', 0) for q in exam.get('quiz', []))
        total_duration += exam_duration
        
        # Create table with modules + exam + total
        table = doc.add_table(rows=1+len(structure.get('modules', []))+2, cols=3)
        table.style = 'Table Grid'
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.8)  # SECTION
            row.cells[1].width = Inches(3.0)  # TITLE
            row.cells[2].width = Inches(1.0)  # DURATION
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'SECTION'
        hdr_cells[1].text = 'TITLE'
        hdr_cells[2].text = 'DURATION'
        
        # Modules
        for i, module in enumerate(structure.get('modules', []), 1):
            module_duration = sum(v.get('duration_minutes', 0) for v in module.get('videos', []))
            module_duration += sum(r.get('duration_minutes', 0) for r in module.get('readings', []))
            module_duration += sum(c.get('duration_minutes', 0) for c in module.get('case_studies', []))
            module_duration += sum(q.get('duration_minutes', 0) for q in module.get('quiz', []))
            
            row_cells = table.rows[i].cells
            row_cells[0].text = f"Module {module.get('module_number', i)}"
            row_cells[1].text = module.get('title', 'Untitled Module')
            row_cells[2].text = f"{module_duration} min"
        
        # Exam
        exam_row = table.rows[len(structure.get('modules', []))+1].cells
        exam_row[0].text = "Final Exam"
        exam_row[1].text = "Final Exam"
        exam_row[2].text = f"{exam_duration} min"
        
        # Total
        total_row = table.rows[-1].cells
        total_row[0].text = "TOTAL"
        total_row[1].text = ""
        total_row[2].text = f"{total_duration} min"

    def add_modules(self, doc, structure):
        """Add modules with all learning elements, durations, and total duration"""
        for module in structure.get('modules', []):
            # Module title
            mod_title = doc.add_paragraph(f"Module {module.get('module_number', 1)}: {module.get('title', 'Untitled Module')}", style='ModuleTitle')
            mod_title.runs[0].font.size = Pt(20)
            
            # Calculate module duration
            module_duration = 0
            module_duration += sum(v.get('duration_minutes', 0) for v in module.get('videos', []))
            module_duration += sum(r.get('duration_minutes', 0) for r in module.get('readings', []))
            module_duration += sum(c.get('duration_minutes', 0) for c in module.get('case_studies', []))
            module_duration += sum(q.get('duration_minutes', 0) for q in module.get('quiz', []))
            
            # Info block
            info_p = doc.add_paragraph()
            info_p.add_run(f"⏱️ Duration: {module_duration} minutes | ").bold = True
            info_p.add_run(f"🎥 {len(module.get('videos', []))} Videos | ").bold = True
            info_p.add_run(f"📖 {len(module.get('readings', []))} Readings | ").bold = True
            info_p.add_run(f"🧠 {len(module.get('case_studies', []))} Case Studies | ").bold = True
            info_p.add_run(f"❓ {len(module.get('quiz', []))} Quiz Questions").bold = True
            
            # Table for learning elements
            elements = []
            elements.extend([{'type': 'Video', 'data': v} for v in module.get('videos', [])])
            elements.extend([{'type': 'Reading', 'data': r} for r in module.get('readings', [])])
            elements.extend([{'type': 'Case Study', 'data': c} for c in module.get('case_studies', [])])
            elements.extend([{'type': 'Quiz', 'data': q} for q in module.get('quiz', [])])
            
            table = doc.add_table(rows=1+len(elements)+1, cols=4)
            table.style = 'Table Grid'
            
            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(0.8)  # TYPE
                row.cells[1].width = Inches(3.0)  # TITLE
                row.cells[2].width = Inches(4.2)  # DESCRIPTION
                row.cells[3].width = Inches(1.0)  # DURATION
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'TYPE'
            hdr_cells[1].text = 'TITLE'
            hdr_cells[2].text = 'DESCRIPTION'
            hdr_cells[3].text = 'DURATION'
            
            # Elements
            for i, elem in enumerate(elements):
                row_cells = table.rows[i+1].cells
                row_cells[0].text = elem['type'].upper()
                row_cells[1].text = elem['data'].get('title', 'QUESTION' if elem['type'] == 'Quiz' else 'Untitled')
                row_cells[3].text = f"{elem['data'].get('duration_minutes', 0)} min"
                
                # Full description without truncation
                desc_text = elem['data'].get('description', 'No description available')
                
                if elem['type'] == 'Case Study':
                    outcomes = elem['data'].get('learning_outcomes', [])
                    desc_text += f"\nLEARNING OUTCOMES:\n" + "\n".join([f"• {o}" for o in outcomes])
                    questions = elem['data'].get('questions', [])
                    question_texts = [q['text'] if isinstance(q, dict) and 'text' in q else str(q) for q in questions]
                    desc_text += f"\nQUESTIONS ({len(question_texts)}):\n" + "\n".join([f"{j+1}. {q}" for j, q in enumerate(question_texts)])
                
                elif elem['type'] == 'Quiz':
                    desc_text = f"{elem['data'].get('question', 'N/A')}\nOPTIONS: {', '.join(elem['data'].get('options', ['N/A']))}\nANSWER: {elem['data'].get('answer', 'N/A')}"
                
                row_cells[2].text = desc_text
            
            # Total duration row
            total_row = table.rows[-1].cells
            total_row[0].text = "TOTAL"
            total_row[1].text = ""
            total_row[2].text = ""
            total_row[3].text = f"{module_duration} min"
            
            doc.add_paragraph()  # Space

    def add_final_exam(self, doc, structure):
        """Add final exam section with duration and total"""
        exam = structure.get('exam', {})
        exam_title = doc.add_heading('FINAL EXAM', level=1)
        exam_title.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Calculate exam duration
        exam_duration = 0
        exam_duration += sum(c.get('duration_minutes', 0) for c in exam.get('case_studies', []))
        exam_duration += sum(q.get('duration_minutes', 0) for q in exam.get('quiz', []))
        
        exam_p = doc.add_paragraph()
        exam_p.add_run(f"⏱️ Duration: {exam_duration} minutes | ").bold = True
        exam_p.add_run(f"✅ Passing Threshold: {exam.get('passing_threshold', 'N/A')}").bold = True
        
        # Table for exam elements
        elements = []
        elements.extend([{'type': 'Case Study', 'data': c} for c in exam.get('case_studies', [])])
        elements.extend([{'type': 'Quiz', 'data': q} for q in exam.get('quiz', [])])
        
        if elements:
            table = doc.add_table(rows=1+len(elements)+1, cols=4)
            table.style = 'Table Grid'
            
            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(0.8)  # TYPE
                row.cells[1].width = Inches(3.0)  # TITLE
                row.cells[2].width = Inches(4.2)  # DESCRIPTION
                row.cells[3].width = Inches(1.0)  # DURATION
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'TYPE'
            hdr_cells[1].text = 'TITLE'
            hdr_cells[2].text = 'DESCRIPTION'
            hdr_cells[3].text = 'DURATION'
            
            # Elements
            for i, elem in enumerate(elements):
                row_cells = table.rows[i+1].cells
                row_cells[0].text = elem['type'].upper()
                row_cells[1].text = elem['data'].get('title', 'QUESTION' if elem['type'] == 'Quiz' else 'Untitled')
                row_cells[3].text = f"{elem['data'].get('duration_minutes', 0)} min"
                
                # Full description without truncation
                desc_text = elem['data'].get('description', 'No description available')
                
                if elem['type'] == 'Case Study':
                    outcomes = elem['data'].get('learning_outcomes', [])
                    desc_text += f"\nLEARNING OUTCOMES:\n" + "\n".join([f"• {o}" for o in outcomes])
                    questions = elem['data'].get('questions', [])
                    question_texts = [q['text'] if isinstance(q, dict) and 'text' in q else str(q) for q in questions]
                    desc_text += f"\nQUESTIONS ({len(question_texts)}):\n" + "\n".join([f"{j+1}. {q}" for j, q in enumerate(question_texts)])
                
                elif elem['type'] == 'Quiz':
                    desc_text = f"{elem['data'].get('question', 'N/A')}\nOPTIONS: {', '.join(elem['data'].get('options', ['N/A']))}\nANSWER: {elem['data'].get('answer', 'N/A')}"
                
                row_cells[2].text = desc_text
            
            # Total duration row
            total_row = table.rows[-1].cells
            total_row[0].text = "TOTAL"
            total_row[1].text = ""
            total_row[2].text = ""
            total_row[3].text = f"{exam_duration} min"