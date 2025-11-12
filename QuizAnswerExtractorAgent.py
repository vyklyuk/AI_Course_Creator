# -*- coding: utf-8 -*-
"""
QuizAnswerExtractorAgent
------------------------
Extracts quiz from JSON, saves:
1. quiz.json
2. quiz_edx.txt  (>>...<< format)
3. quiz_pretty.docx (with Correct + checkmark)

No Cyrillic in console output.
"""

import json
import os
from typing import Any, Dict, List
from datetime import datetime

# === DOCX IMPORTS ===
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Optional, Tuple


class QuizAnswerExtractorAgent:
    def __init__(self, verbose = True):
        self.verbose = verbose

    # ==================================================================== #
    # 1. FIND QUIZ + PATH
    # ==================================================================== #
    def _find_quiz_and_path(self, data: Any) -> tuple[List, str]:
        if isinstance(data, dict):
            if "questions" in data and data.get("path"):
                return data.get("questions", []), data["path"]
            for key, value in data.items():
                if key == "quiz" and isinstance(value, dict) and value.get("path"):
                    return value.get("questions", []), value["path"]
                q, p = self._find_quiz_and_path(value)
                if q is not None:
                    return q, p
        elif isinstance(data, list):
            return data, "output/unknown_quiz"
        return None, None

    # ==================================================================== #
    # 2. BUILD QUESTION (for JSON)
    # ==================================================================== #
    def _build_question(self, raw: Dict, idx: int) -> Dict:
        q_text = raw.get("question", f"Q{idx}")
        options = raw.get("options", [])
        answer = raw.get("answer")

        struct = {
            "id": idx,
            "question": q_text,
            "options": options,
            "correct_answer": answer
        }
        if options and answer is not None:
            try:
                struct["correct_index"] = options.index(answer)
            except ValueError:
                struct["correct_index"] = None
        return struct

    # ==================================================================== #
    # 3. SAVE JSON + TRIGGER OTHER FORMATS
    # ==================================================================== #
    def extract_and_save(self, json_data: Any) -> Dict:
        if isinstance(json_data, str):
            json_data = json.loads(json_data)

        questions_raw, save_folder = self._find_quiz_and_path(json_data)
        if not questions_raw:
            raise ValueError("No questions found")
        if not save_folder:
            raise ValueError("Missing 'path' in quiz object")

        os.makedirs(save_folder, exist_ok=True)
        json_path = os.path.join(save_folder, "quiz.json")

        questions = [
            self._build_question(q, i + 1)
            for i, q in enumerate(questions_raw)
            if isinstance(q, dict)
        ]

        result = {
            "total_questions": len(questions),
            "extracted_at": datetime.now().isoformat(),
            "source_path": save_folder,
            "questions": questions,
            "path": os.path.abspath(json_path)
        }

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        if self.verbose:
            print(f"Saved JSON: {result['path']}")

        # Save other formats
        self.save_as_edx_txt(json_data)
        self.save_as_docx(json_data)

        return result

    # ==================================================================== #
    # 4. SAVE eDx TXT (>>...<<)
    # ==================================================================== #
    def save_as_edx_txt(self, json_data: Any) -> str:
        if isinstance(json_data, str):
            json_data = json.loads(json_data)

        questions_raw, save_folder = self._find_quiz_and_path(json_data)
        if not questions_raw or not save_folder:
            raise ValueError("No data or path")

        os.makedirs(save_folder, exist_ok=True)
        lines = []

        for q in questions_raw:
            if not isinstance(q, dict):
                continue
            prompt = q.get("question", "").strip()
            options = q.get("options", [])
            correct = q.get("answer")

            lines.append(f">>{prompt}<<")
            lines.append("")

            for opt in options:
                marker = "(x)" if opt == correct else "( )"
                lines.append(f"{marker} {opt}")

            lines.append("")

        content = "\n".join(lines).rstrip()
        file_path = os.path.join(save_folder, "quiz_edx.txt")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        abs_path = os.path.abspath(file_path)
        if self.verbose:
            print(f"Saved eDx TXT: {abs_path}")
        return abs_path

    # ==================================================================== #
    # 5. SAVE PRETTY DOCX
    # ==================================================================== #
    def save_as_docx(self, json_data: Any) -> str:
        if isinstance(json_data, str):
            json_data = json.loads(json_data)

        questions_raw, save_folder = self._find_quiz_and_path(json_data)
        if not questions_raw or not save_folder:
            raise ValueError("No data or path")

        os.makedirs(save_folder, exist_ok=True)
        doc = Document()
        self._setup_doc_styles(doc)

        # Title
        title = doc.add_heading("Test Module", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Questions
        for i, q in enumerate(questions_raw, 1):
            if not isinstance(q, dict):
                continue

            question_text = q.get("question", "Question missing")
            options = q.get("options", [])
            correct = q.get("answer")

            # Question number
            p_num = doc.add_paragraph()
            run = p_num.add_run(f"Question {i}")
            run.bold = True
            run.font.size = Pt(13)

            # Question text
            p_q = doc.add_paragraph(question_text, style='Intense Quote')
            p_q.paragraph_format.space_after = Pt(6)

            # Options
            for j, opt in enumerate(options, 65):
                p_opt = doc.add_paragraph()
                run = p_opt.add_run(f"  {chr(j)}. {opt}")
                run.font.size = Pt(11)

                if opt == correct:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 0)
                    check_run = p_opt.add_run(" Correct")
                    check_run.font.color.rgb = RGBColor(0, 140, 0)
                    check_run.bold = True
                    self._add_checkmark(p_opt)
                else:
                    run.font.color.rgb = RGBColor(70, 70, 70)

            doc.add_paragraph()

        # Save
        file_path = os.path.join(save_folder, "quiz_pretty.docx")
        doc.save(file_path)
        abs_path = os.path.abspath(file_path)
        if self.verbose:
            print(f"Saved pretty DOCX: {abs_path}")
        return abs_path

    # ==================================================================== #
    # HELPER: STYLES & CHECKMARK
    # ==================================================================== #
    def _setup_doc_styles(self, doc):
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        style = doc.styles['Heading 1']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(16)
        font.bold = True

        style = doc.styles['Intense Quote']
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Pt(20)

    def _add_checkmark(self, paragraph):
        """Adds green checkmark via Wingdings (no text)"""
        run = paragraph.add_run()
        run.text = "P"  # 'P' in Wingdings = checkmark
        run.font.name = 'Wingdings'
        run.font.color.rgb = RGBColor(0, 140, 0)
        run.font.size = Pt(14)


# ==================================================================== #
# EXAMPLE USAGE
# ==================================================================== #
if __name__ == "__main__":
    with open("AML_Uncovered_curriculum_with_paths.json", "r", encoding="utf-8") as f:
        data = json.load(f)

    agent = QuizAnswerExtractorAgent()
    module2 = data["modules"][1]
    agent.extract_and_save(module2)